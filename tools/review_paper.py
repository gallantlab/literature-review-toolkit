#!/usr/bin/env python3
"""Phase 7 — build a review ARTICLE (.docx) from a finished review corpus.

This tool owns only the *mechanics*: the title/author/disclosure block, the
abstract, the section headings + body paragraphs, an embedded figure with a
standalone caption, and an APA-7 **reference list pulled straight from
rows.json** (the canonical `apa` strings, deduped, alphabetised, hanging-indent).

It does NOT write prose. Author the prose separately with the scientific-writing
skill and supply it as a content JSON (`--content`). In-text citations are APA
author-date, e.g. "(Huth et al., 2016)"; every in-text citation must name a paper
that exists in rows.json so the reference list backs it.

If the review is AI-authored, say so: put the model's name in `authors`, an
`author_note` identifying it as an AI, and a `disclosure` paragraph (the
fabricate-then-verify caveat — see PLAYBOOK Phase 7).

content.json schema:
{
  "title": "single string (\\n allowed for a two-line title)",
  "authors": ["Claude Fable 5"],
  "author_note": "An artificial intelligence developed by Anthropic",
  "affiliation_line": "Prepared for ... · DD Month YYYY",
  "disclosure": "Author's disclosure: ...",          # optional
  "abstract": "one paragraph",
  "sections": [ {"heading": "1. Introduction", "level": 1,
                 "paragraphs": ["para one", "para two"]}, ... ],
  "figure": {"path": "<topic>_families.png", "caption": "Figure 1. ..."},  # optional; --figure overrides
  "references_heading": "References",                 # optional (default "References")
  "references_note": "..."                            # optional; "{n}" is replaced by the count
}

Usage:
  python3 tools/review_paper.py --rows rows.json --content content.json \
          --out <Topic>_review.docx [--figure <topic>_families.png]
"""
import json, re, argparse, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _apa_sort_key(apa):
    return re.sub(r"[^A-Za-z]", "", apa).lower()


def _fmt_authors(authors):
    authors = [a for a in authors if a]
    if len(authors) <= 1:
        return authors[0] if authors else ""
    if len(authors) == 2:
        return f"{authors[0]} & {authors[1]}"
    return ", ".join(authors[:-1]) + f", & {authors[-1]}"


def build(content, rows, out, figure=None):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    def centered(text, size, bold=False, italic=False, color=None):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color: r.font.color.rgb = RGBColor(*color)
        return p

    # ---- title / author / disclosure block
    centered(content["title"], 17, bold=True)
    if content.get("authors"):
        centered(_fmt_authors(content["authors"]), 12, bold=True)
    if content.get("author_note"):
        centered(content["author_note"], 10.5, italic=True)
    if content.get("affiliation_line"):
        centered(content["affiliation_line"], 9.5, color=(0x66, 0x66, 0x66))
    if content.get("disclosure"):
        centered(content["disclosure"], 8.5, italic=True, color=(0x88, 0x88, 0x88))

    # ---- abstract
    if content.get("abstract"):
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(content["abstract"])

    # ---- body sections
    for sec in content.get("sections", []):
        doc.add_heading(sec["heading"], level=sec.get("level", 1))
        for para in sec.get("paragraphs", []):
            doc.add_paragraph(para)

    # ---- figure (CLI --figure overrides content.figure.path)
    fig = content.get("figure") or {}
    fig_path = figure or fig.get("path")
    if fig_path:
        if not os.path.isabs(fig_path):
            fig_path = os.path.join(os.path.dirname(os.path.abspath(out)) or ".", fig_path)
        doc.add_page_break()
        doc.add_heading("Figure", level=1)
        fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run().add_picture(fig_path, width=Inches(6.4))
        if fig.get("caption"):
            cap = doc.add_paragraph()
            cap.add_run(fig["caption"])
            for r in cap.runs:
                r.font.size = Pt(9)

    # ---- references: APA-7 from rows.json apa, deduped + alphabetised + hanging indent
    doc.add_page_break()
    doc.add_heading(content.get("references_heading", "References"), level=1)
    seen, refs = set(), []
    for row in rows:
        apa = (row.get("apa") or "").strip()
        if apa and apa not in seen:
            seen.add(apa)
            refs.append((apa, row.get("link", "")))
    refs.sort(key=lambda x: _apa_sort_key(x[0]))
    for apa, link in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)   # APA hanging indent
        p.paragraph_format.space_after = Pt(4)
        p.add_run(apa if apa.endswith(".") else apa + ".")
        if link:
            p.add_run(" " + link).font.color.rgb = RGBColor(0x40, 0x40, 0x80)
        for r in p.runs:
            r.font.size = Pt(9.5)
    if content.get("references_note"):
        np = doc.add_paragraph()
        nr = np.add_run(content["references_note"].replace("{n}", str(len(refs))))
        nr.italic = True; nr.font.size = Pt(9)

    doc.save(out)
    return len(refs)


def main():
    ap = argparse.ArgumentParser(description="Phase 7: build a review-article .docx from a corpus.")
    ap.add_argument("--rows", required=True, help="rows.json (canonical apa supplies the reference list)")
    ap.add_argument("--content", required=True, help="content.json (prose: title/author/abstract/sections/figure)")
    ap.add_argument("--out", required=True, help="output .docx path")
    ap.add_argument("--figure", help="figure image path (overrides content.figure.path)")
    args = ap.parse_args()
    rows = json.load(open(args.rows, encoding="utf-8"))
    content = json.load(open(args.content, encoding="utf-8"))
    n = build(content, rows, args.out, figure=args.figure)
    print(f"wrote {args.out} | {n} references (APA-7, from {args.rows})")


if __name__ == "__main__":
    main()
